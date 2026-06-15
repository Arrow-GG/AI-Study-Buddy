"""Document ingestion and retrieval helpers."""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from fastapi import HTTPException, UploadFile, status
from sqlalchemy.orm import Session

from app.config import settings
from app.database import Document, DocumentChunk, User, utcnow
from app.services.ai import add_embeddings_to_chunks


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def normalize_id(value: int | str) -> int:
    try:
        return int(value)
    except (TypeError, ValueError) as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid id") from exc


def serialize_document(document: Document) -> dict:
    return {
        "id": document.id,
        "filename": document.filename,
        "file_path": document.file_path,
        "file_size": document.file_size,
        "mime_type": document.mime_type,
        "status": document.status,
        "upload_date": document.upload_date.isoformat(),
        "processed_date": document.processed_date.isoformat() if document.processed_date else None,
        "text_length": document.text_length,
    }


def get_document_or_404(db: Session, user: User, document_id: int | str) -> Document:
    document = (
        db.query(Document)
        .filter(Document.id == normalize_id(document_id), Document.user_id == user.id)
        .first()
    )
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return document


def validate_file(file: UploadFile) -> None:
    content_type = file.content_type or "application/octet-stream"
    suffix = Path(file.filename or "").suffix.lower()
    allowed_suffixes = {".pdf", ".docx", ".txt"}
    if content_type not in settings.ALLOWED_FILE_TYPES and suffix not in allowed_suffixes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF, DOCX, and TXT files are supported",
        )


async def save_upload(db: Session, user: User, file: UploadFile) -> dict:
    validate_file(file)
    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)

    document = Document(
        user_id=user.id,
        filename=Path(file.filename or "document.txt").name,
        file_path="",
        file_size=0,
        mime_type=file.content_type or "application/octet-stream",
        status="processing",
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    original_name = document.filename
    stored_name = f"{user.id}-{document.id}-{original_name}"
    file_path = upload_dir / stored_name

    size = 0
    with file_path.open("wb") as handle:
        while chunk := await file.read(1024 * 1024):
            size += len(chunk)
            if size > settings.MAX_FILE_SIZE:
                file_path.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail="File exceeds the maximum upload size",
                )
            handle.write(chunk)

    document.file_path = str(file_path)
    document.file_size = size
    db.commit()
    db.refresh(document)
    return serialize_document(process_document(db, user, document.id))


def save_text_document(db: Session, user: User, title: str, content: str) -> dict:
    cleaned = content.strip()
    if len(cleaned) < 20:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Paste at least 20 characters of article or note text",
        )

    encoded = cleaned.encode("utf-8")
    if len(encoded) > settings.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="Text exceeds the maximum document size",
        )

    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)

    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "-", title.strip() or "pasted-notes").strip("-")
    filename = f"{safe_title[:80] or 'pasted-notes'}.txt"
    document = Document(
        user_id=user.id,
        filename=filename,
        file_path="",
        file_size=len(encoded),
        mime_type="text/plain",
        status="processing",
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    stored_name = f"{user.id}-{document.id}-{filename}"
    file_path = upload_dir / stored_name
    file_path.write_text(cleaned, encoding="utf-8")
    document.file_path = str(file_path)
    db.commit()
    db.refresh(document)
    return serialize_document(process_document(db, user, document.id))


def process_document(db: Session, user: User, document_id: int | str) -> Document:
    document = get_document_or_404(db, user, document_id)
    try:
        text = extract_text(Path(document.file_path), document.mime_type)
        chunks = chunk_text(text)
        db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()
        chunk_models = []
        for index, content in enumerate(chunks):
            chunk_models.append(
                DocumentChunk(
                    document_id=document.id,
                    content=content,
                    chunk_index=index,
                    chunk_metadata={
                        "filename": document.filename,
                        "chunk": index + 1,
                        "document_id": document.id,
                    },
                )
            )
        add_embeddings_to_chunks(chunk_models)
        db.add_all(chunk_models)
        document.status = "completed"
        document.processed_date = utcnow()
        document.text_length = len(text)
        document.error = None
        db.commit()
        db.refresh(document)
        return document
    except Exception as exc:
        document.status = "failed"
        document.error = str(exc)
        db.commit()
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not process document: {exc}",
        ) from exc


def extract_text(path: Path, mime_type: str) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt" or mime_type == "text/plain":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf" or mime_type == "application/pdf":
        return extract_pdf_text(path)
    if suffix == ".docx" or "wordprocessingml" in mime_type:
        return extract_docx_text(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf to be installed") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx to be installed") from exc

    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def chunk_text(text: str) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        raise RuntimeError("No readable text found in the document")

    chunk_size = settings.CHUNK_SIZE
    overlap = min(settings.CHUNK_OVERLAP, chunk_size // 2)
    chunks = []
    start = 0
    while start < len(cleaned):
        end = min(len(cleaned), start + chunk_size)
        chunks.append(cleaned[start:end].strip())
        if end == len(cleaned):
            break
        start = max(0, end - overlap)
    return chunks


def delete_document(db: Session, user: User, document_id: int | str) -> None:
    document = get_document_or_404(db, user, document_id)
    Path(document.file_path).unlink(missing_ok=True)
    db.delete(document)
    db.commit()


def document_chunks(db: Session, user: User, document_id: int | str) -> list[dict]:
    document = get_document_or_404(db, user, document_id)
    chunks = (
        db.query(DocumentChunk)
        .filter(DocumentChunk.document_id == document.id)
        .order_by(DocumentChunk.chunk_index)
        .all()
    )
    return [
        {
            "id": chunk.id,
            "document_id": chunk.document_id,
            "content": chunk.content,
            "chunk_index": chunk.chunk_index,
            "metadata": chunk.chunk_metadata or {},
            "created_at": chunk.created_at.isoformat(),
        }
        for chunk in chunks
    ]


def all_document_text(db: Session, user: User, document_id: int | str) -> str:
    return "\n\n".join(chunk["content"] for chunk in document_chunks(db, user, document_id))
